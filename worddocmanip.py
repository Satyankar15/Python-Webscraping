# -*- coding: utf-8 -*-
"""
Created on Thu Oct  8 09:51:38 2020

@author: satya
"""

import docx
d=docx.Document(r"C:\Users\satya\Desktop\Desktop Documents/DISADVANTAGES of VSLAM.docx")
#print(d.paragraphs)
p=d.paragraphs[0]
x=5
print(p.runs[0].underline)
p.runs[1].italic= True
print(p.text)
d.save(r"C:\Users\satya\Desktop\Desktop Documents/DISADVANTAGES of VSLAM2.docx")
x=x+1